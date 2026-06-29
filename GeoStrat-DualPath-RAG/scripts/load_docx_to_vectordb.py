"""
将福建地质志.docx存入向量库（FAISS）
"""
import os, sys, re, time
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import docx
from langchain_core.documents import Document
from rag_core import vectordb


def extract_and_chunk(docx_path, min_chars=200, max_chars=1200):
    doc = docx.Document(docx_path)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text or re.match(r'^\s*\d{1,4}\s*$', text):
            continue
        paragraphs.append(text)

    # 合并短段落，切分长段落
    result = []
    buffer = ""
    for t in paragraphs:
        if len(t) >= max_chars:
            if buffer:
                result.append(buffer.strip()); buffer = ""
            for s in re.split(r'(?<=[。！？])', t):
                if len(buffer) + len(s) > max_chars and buffer:
                    result.append(buffer.strip()); buffer = s
                else:
                    buffer += s
            if buffer.strip():
                result.append(buffer.strip()); buffer = ""
        elif len(t) < min_chars:
            buffer = (buffer + "\n" + t).strip() if buffer else t
            if len(buffer) >= min_chars:
                result.append(buffer.strip()); buffer = ""
        else:
            if buffer:
                result.append(buffer.strip()); buffer = ""
            result.append(t)
    if buffer.strip():
        result.append(buffer.strip())
    return [c for c in result if 150 <= len(c) <= 2000]


def main():
    chunks = extract_and_chunk(r'E:\地质文本空间化平台\福建地质志.docx')
    print(f'{len(chunks)} chunks')

    batch_size = 50
    for i in range(0, len(chunks), batch_size):
        batch = chunks[i:i + batch_size]
        docs = [Document(page_content=c, metadata={
            'source_file': '福建地质志.docx',
            'kb_id': 'experiment',
            'chunk_index': i + j,
            'char_count': len(c)
        }) for j, c in enumerate(batch)]
        vectordb.add_documents(docs)
        pct = min(i + batch_size, len(chunks))
        if pct % 200 == 0 or pct == len(chunks):
            print(f'  [{pct}/{len(chunks)}]')

    print(f'Done: {vectordb._index.ntotal} vectors stored')


if __name__ == '__main__':
    main()
